import streamlit as st
import random
import sqlite3
from docx import Document
from io import BytesIO

# Establish connection and create table if it doesn't exist
conn = sqlite3.connect("gift_exchange_uiuc.db")
c = conn.cursor()

# Ensure users table exists
c.execute('''
    CREATE TABLE IF NOT EXISTS users (
        name TEXT PRIMARY KEY,
        password TEXT,
        gift TEXT
    )
''')
conn.commit()

# Updated participants list
participants = ["Emily", "Rebecca", "Kelly", "Joanne", "Cecilie", "Sammy", "Cathy", "Woody", "Jensen", "Ali", "Joseph", "Steven", "Alan", "Angelina", "Ericka", "Paul"]
total_participants = len(participants)  # Total number of participants

# Shuffle participants and assign gifts ensuring no one draws themselves
def assign_gifts(participants):
    receivers = participants.copy()
    random.shuffle(receivers)
    while any(p == r for p, r in zip(participants, receivers)):
        random.shuffle(receivers)
    return dict(zip(participants, receivers))

# Initialize session state for gift assignments and page
if "gift_assignments" not in st.session_state:
    st.session_state["gift_assignments"] = assign_gifts(participants)

if "page" not in st.session_state:
    st.session_state["page"] = "login_or_setup"

# Function to count the number of participants who have drawn
def count_completed_draws():
    c.execute("SELECT COUNT(*) FROM users WHERE gift IS NOT NULL")
    result = c.fetchone()
    return result[0] if result else 0

# Function to create a Word document with pairing results
def create_pairing_doc():
    document = Document()
    document.add_heading("交換禮物配對結果", level=1)

    c.execute("SELECT name, gift FROM users")
    results = c.fetchall()
    for row in results:
        giver = row[0]
        receiver = row[1]
        document.add_paragraph(f"{giver} 抽到的對象是：{receiver}")

    # Save document to BytesIO for downloading
    doc_io = BytesIO()
    document.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Admin view to see all gift assignments and download Word file
def admin_view():
    st.title("管理員模式 - 查看所有抽籤結果")
    admin_password = st.text_input("請輸入管理員密碼", type="password")

    if st.button("登入管理員模式"):
        if admin_password == "admin123":
            st.write("管理員登入成功！")

            # Check if all participants have completed the draw
            completed_draws = count_completed_draws()
            if completed_draws == total_participants:
                # Download button for pairing results if everyone has drawn
                doc_io_pairing = create_pairing_doc()
                st.download_button(
                    label="下載配對結果 Word 檔案",
                    data=doc_io_pairing,
                    file_name="pairing_results.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.write(f"目前還有 {total_participants - completed_draws} 人尚未完成抽籤，無法下載配對結果。")
        else:
            st.write("管理員密碼錯誤！")

# Login or set password page with instructions
if st.session_state["page"] == "login_or_setup":
    st.title("交換禮物一夜暴富🤑")

    # Instructions
    st.write("🎄Merry Pre-Christmas🎄")
    st.write("如果沒有創建過帳號，請先輸入英文名字（第一個字母大寫如 Justin），並輸入一組密碼。")
    st.write("先按「設定密碼」，再按「登入」。（如果沒有反應就多按一次按鈕哈哈哈）")

    name = st.text_input("請輸入你的名字")
    password = st.text_input("請輸入密碼", type="password")

    if st.button("設定密碼"):
        if name in participants:
            c.execute("SELECT * FROM users WHERE name = ?", (name,))
            user = c.fetchone()
            if user is None:
                c.execute("INSERT INTO users (name, password) VALUES (?, ?)", (name, password))
                conn.commit()
                st.write("密碼設定成功！")
                st.session_state["page"] = "draw"
                st.session_state["current_user"] = name
            else:
                st.write("此使用者已設定過密碼，請直接登入。")
        else:
            st.write("這個名字不在參與者名單中，請確認後再試。")

    if st.button("登入"):
        c.execute("SELECT * FROM users WHERE name = ? AND password = ?", (name, password))
        user = c.fetchone()
        if user:
            st.write("登入成功！")
            st.session_state["current_user"] = name
            if user[2]:
                st.session_state["page"] = "view"
            else:
                st.session_state["page"] = "draw"
        else:
            st.write("名字或密碼錯誤，請重新輸入。")

    if st.button("管理員登入"):
        st.session_state["page"] = "admin"

# Draw page with instructions
elif st.session_state["page"] == "draw":
    st.title("交換禮物一夜暴富🤑 - 抽籤")
    name = st.session_state["current_user"]

    # Instructions for the draw page
    st.write("請按下「抽籤」按鈕就會看到你要當誰的聖誕老人HOHOHO")
    st.write("再按一次抽籤會到結果頁面，如果忘記自己抽到誰可以再次登入查看")
    st.write("🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨")
    st.write(" ")

    # Check if user already has a gift
    c.execute("SELECT gift FROM users WHERE name = ?", (name,))
    result = c.fetchone()
    if result and result[0]:
        # If the user has already drawn, show only the result and "Return to Main Page" button
        drawn_gift = result[0]
        st.write(f"🎉 {name}，你是 {drawn_gift} 的聖誕老人 🎉")
        if st.button("返回主頁"):
            st.session_state["page"] = "login_or_setup"
    else:
        # If the user has not drawn yet, show the "Draw" button and remaining count
        completed_count = count_completed_draws()
        remaining_count = total_participants - completed_count
        st.write(f"目前尚有 {remaining_count} 人尚未抽籤")
        if st.button("抽籤"):
            gift = st.session_state["gift_assignments"][name]
            c.execute("UPDATE users SET gift = ? WHERE name = ?", (gift, name))
            conn.commit()
            st.write(f"🎉 {name}，你是 {gift} 的聖誕老人 🎉")
            # Change the page to "view" and let Streamlit rerender
            st.session_state["page"] = "view"
            st.session_state["rerun_flag"] = not st.session_state.get("rerun_flag", False)  # Toggle rerun flag

# View result page with only "Return to Main Page" button
elif st.session_state["page"] == "view":
    st.title("交換禮物一夜暴富🤑 - 查看結果")
    name = st.session_state["current_user"]

    c.execute("SELECT gift FROM users WHERE name = ?", (name,))
    result = c.fetchone()
    if result and result[0]:
        st.write("🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨🎄🎅🏻🎁🦌✨")
        st.write(f"{name}，你是 {result[0]} 的聖誕老人")
    else:
        st.write("你尚未抽籤。")

    if st.button("返回主頁"):
        st.session_state["page"] = "login_or_setup"

# Admin page view
elif st.session_state["page"] == "admin":
    admin_view()
    if st.button("返回主頁"):
        st.session_state["page"] = "login_or_setup"
